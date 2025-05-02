import streamlit as st
import pandas as pd
import numpy as np
import io
import plotly.express as px
import plotly.graph_objects as go
import time
import re # Import regex for salary extraction
from collections import Counter # Import Counter for missing skills
import os # Import os if needed for file paths (though not strictly used here)

# Import your custom modules
# Make sure these modules are in your project directory and updated
try:
    # Import the updated parser function for structure analysis
    from resume_parser import extract_text_from_pdf, extract_text_from_docx, analyze_resume as analyze_resume_structure
    # Import functions from the updated semantic matcher
    from semantic_matcher import (
        semantic_match_resume,
        extract_resume_keywords, # Keep if you use it elsewhere, though not in the main analysis tab now
        extract_skills_from_resume,
        get_skill_matches,
        # The analyze_resume function was removed from semantic_matcher
    )
    # Import visualization libraries needed directly in app.py
    from wordcloud import WordCloud
    import matplotlib.pyplot as plt
except ImportError as e:
    st.error(f"Missing required modules or functions: {e}. Please ensure resume_parser.py, semantic_matcher.py, and necessary libraries are installed and up to date.")
    st.stop() # Stop execution if modules is missing

# Page configuration with custom theme
st.set_page_config(
    page_title="Smart Job Matcher",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "# Smart Job Matcher\nFind the perfect job match for your skills and experience!"
    }
)

# Apply custom styling
# Added some basic styles for the custom classes used in markdown
st.markdown("""
<style>
    .main-header {
        font-size: 2.5em;
        font-weight: bold;
        color: #1E3A8A; /* Dark Blue */
        text-align: center;
        margin-bottom: 30px;
        text-shadow: 1px 1px 2px #cccccc;
    }
    .sub-header {
        font-size: 1.8em;
        font-weight: bold;
        color: #3B82F6; /* Medium Blue */
        margin-top: 20px;
        margin-bottom: 15px;
        border-bottom: 2px solid #EFF6FF; /* Light Blue */
        padding-bottom: 5px;
    }
    .match-section {
        border: 1px solid #D1D5DB; /* Light Gray */
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 20px;
        box-shadow: 2px 2px 8px rgba(0,0,0,0.1);
    }
    .job-title {
        font-size: 1.3em;
        font-weight: bold;
        color: #1E40AF; /* Darker Blue */
        margin-bottom: 5px;
    }
    .match-score-high {
        color: #16A34A; /* Green */
        font-weight: bold;
    }
    .match-score-medium {
        color: #D97706; /* Amber */
        font-weight: bold;
    }
    .match-score-low {
        color: #DC2626; /* Red */
        font-weight: bold;
    }
    .matched-keywords {
        color: #065F46; /* Dark Green */
        font-size: 0.9em;
        margin-top: 10px;
    }
    .missing-keywords {
        color: #991B1B; /* Dark Red */
        font-size: 0.9em;
        margin-top: 5px;
    }
     .tips-section {
        background-color: #F3F4F6; /* Lighter Gray */
        border-left: 4px solid #60A5FA; /* Blue */
        padding: 15px;
        border-radius: 4px;
        margin-top: 20px;
        margin-bottom: 20px;
     }
</style>
""", unsafe_allow_html=True)


# Initialize session state for storing resume data
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'resume_structure_analysis' not in st.session_state: # Renamed
    st.session_state.resume_structure_analysis = None
if 'job_results' not in st.session_state:
    st.session_state.job_results = None
if 'skills_extracted' not in st.session_state:
    st.session_state.skills_extracted = []
if 'resume_skills_analysis' not in st.session_state: # Added for skills analysis results
    st.session_state.resume_skills_analysis = None
if 'last_matching_upload_name' not in st.session_state: # To track file changes
    st.session_state.last_matching_upload_name = None
if 'last_analysis_upload_name' not in st.session_state: # To track file changes
    st.session_state.last_analysis_upload_name = None


# Custom header
st.markdown('<div class="main-header">💼 Smart Job Matcher</div>', unsafe_allow_html=True)

# Load job listings with caching
@st.cache_data(ttl=3600)  # Cache data for 1 hour
def load_jobs():
    try:
        # Use an absolute path or path relative to the script location if needed
        # For Streamlit Cloud, files in the same directory are usually accessible directly
        df = pd.read_csv("data/zangia_filtered_jobs.csv")
        # Add additional preprocessing
        df['Salary'] = df['Salary'].fillna('Not specified')
        df['Job description'] = df['Job description'].fillna('')
        df['Requirements'] = df['Requirements'].fillna('')
        # Ensure URL column exists and is string type
        if 'URL' not in df.columns:
            df['URL'] = '#' # Default or placeholder if no URL column
        df['URL'] = df['URL'].astype(str)
        return df
    except FileNotFoundError:
        st.error("Error: zangia_filtered_jobs.csv not found. Please ensure the data file is in a 'data' directory relative to your script.")
        return pd.DataFrame(columns=['Job title', 'Company', 'Salary', 'Job description', 'Requirements', 'URL'])
    except Exception as e:
        st.error(f"Error loading job data: {str(e)}")
        # Return empty dataframe with expected columns
        return pd.DataFrame(columns=['Job title', 'Company', 'Salary', 'Job description', 'Requirements', 'URL'])

# Load jobs data upfront
jobs_df = load_jobs()
if jobs_df.empty:
     # If jobs_df is empty due to loading error, disable job-related features
     job_data_loaded = False
     st.warning("Job data could not be loaded. Job matching and Job Market Explorer are unavailable.")
else:
    job_data_loaded = True


# Grouped sector keywords with enhanced categories
sector_options = {
    "Education & Management": ["багш", "сургалт", "удирдлага", "education", "teacher", "lecturer", "professor", "academic", "сургууль", "их сургууль", "коллеж", "менежмент", "дарга", "захирал"],
    "Customer Service": ["customer service", "үйлчилгээ", "захиалга", "client", "support", "help desk", "call center", "харилцагч", "оператор"],
    "Leadership": ["менежер", "захирал", "manager", "director", "executive", "chief", "head of", "supervisor", "ахлагч", "удирдагч"],
    "Tech & Development": ["developer", "инженер", "програм", "software", "programmer", "coder", "IT", "tech", "web", "mobile", "систем", "сүлжээ", "мэдээллийн технологи", "программист"],
    "Creative & Marketing": ["дизайн", "сошиал", "контент", "media", "designer", "creative", "marketing", "brand", "SEO", "content", "маркетинг", "креатив", "зурагчин", "редактор"],
    "Finance": ["нягтлан", "санхүү", "finance", "accountant", "accounting", "financial", "budget", "tax", "banking", "касс", "данс", "аудит"],
    "Healthcare": ["эмч", "сувилагч", "health", "эрүүл мэнд", "doctor", "nurse", "medical", "healthcare", "clinic", "hospital", "фармаци", "жолооч", "үйлчлэгч"], # Added some potentially related lower-skill roles often in healthcare
    "Logistics & Support": ["logistics", "тээвэр", "нярав", "туслах", "warehouse", "shipping", "supply chain", "inventory", "хангамж", "ложистик", "жолооч", "үйлчлэгч", "ачигч", "мастер"], # Added some potentially related lower-skill roles
    "Data & AI": ["data analyst", "data scientist", "AI", "мэдээлэл шинжээч", "machine learning", "analytics", "big data", "statistics", "шинжилгээ", "дата"],
    "HR & Recruitment": ["HR", "хүний нөөц", "recruiter", "talent", "hiring", "recruitment", "personnel", "staffing"],
    "Legal & Compliance": ["хууль", "lawyer", "legal", "compliance", "attorney", "law", "regulatory", "contracts", "өмгөөлөгч", "эрх зүйч", "комплаенс"],
    "Sales": ["борлуулалт", "sales", "зөвлөх", "business development", "account manager", "retail", "revenue", "худалдаа", "төлөөлөгч"],
    "Project Management": ["төслийн менежер", "project manager", "project coordinator", "scrum", "agile", "program manager", "төсөл", "зохицуулагч"],
    "Engineering & Construction": ["механик", "цахилгаан", "civil", "барилга", "engineer", "mechanical", "electrical", "construction", "инженер", "барилгачин", "цахилгаанчин", "сантехникч", "холбоочин", "баригч"] # Added more specific roles
}
# Create flat keyword list from the grouped options (used in Job Market Explorer filtering)
all_keywords_flat = sorted(set([kw for group in sector_options.values() for kw in group]))


# Sidebar for filters and settings
with st.sidebar:
    st.header("⚙️ Options")
    # App mode selection
    app_mode = st.radio("Select Mode", ["Resume-to-Job Matching", "Resume Analysis", "Job Market Explorer", "Resume Creator"])

    # Add filters specific to Job Market Explorer mode
    if app_mode == "Job Market Explorer" and job_data_loaded:
        st.subheader("Job Market Filters")
        # Use the loaded jobs_df for filter options
        available_companies = sorted(jobs_df['Company'].unique().tolist())
        selected_companies = st.multiselect("Filter by Company", options=available_companies, placeholder="Select companies...", key="explorer_company_filter")

        available_salaries = jobs_df['Salary'].unique().tolist()
        # Try to find min/max numeric salaries for slider if possible
        numeric_salaries = []
        for s in available_salaries:
             if isinstance(s, (int, float)):
                  numeric_salaries.append(s)
             elif isinstance(s, str):
                  # Simple extraction attempt
                  numbers = re.findall(r'\d+', s.replace(',', '')) # Remove commas for large numbers
                  if numbers:
                       # Use the first number found as a proxy
                       numeric_salaries.append(int(numbers[0]))

        min_salary_val = min(numeric_salaries) if numeric_salaries else 0
        # Set a reasonable upper bound if data is missing or max is very high
        max_salary_val = max(numeric_salaries) if numeric_salaries else 10_000_000
        # Ensure max_salary_val is at least min_salary_val if only one job exists etc.
        max_salary_val = max(max_salary_val, min_salary_val + 100000)


        salary_min = st.slider(
            "Minimum Salary (\u20ae)",
            min_value=min_salary_val,
            max_value=max_salary_val,
            value=min_salary_val,
            step=100000,
            help="Filter jobs by minimum advertised salary. Note: Many jobs do not specify salary.",
            key="explorer_salary_filter"
        )

        available_sectors = list(sector_options.keys())
        selected_job_sectors = st.multiselect("Filter by Sector", options=available_sectors, placeholder="Select sectors...", key="explorer_sector_filter")
    elif app_mode == "Job Market Explorer" and not job_data_loaded:
        pass # Message already shown above


    st.markdown("---")
    st.write("#### About")
    st.write("""
    Smart Job Matcher helps you find jobs that match your skills and experience using advanced AI matching technology.
    """)
    st.write("Upload your resume in 'Resume-to-Job Matching' or 'Resume Analysis' modes.")


# === App Modes ===

# === Resume Creator Mode ===
if app_mode == "Resume Creator":
    st.markdown('📄 Resume Creator', unsafe_allow_html=True) # Corrected emoji

    st.info("Fill in the fields below to generate a basic resume (DOCX format)")

    name = st.text_input("Full Name", "", key="creator_name")
    email = st.text_input("Email", "", key="creator_email")
    phone = st.text_input("Phone", "", key="creator_phone")
    linkedin = st.text_input("LinkedIn URL", "", key="creator_linkedin")
    summary = st.text_area("Professional Summary", height=100, value="", key="creator_summary")

    st.markdown("### 🎓 Education") # Corrected emoji
    education = st.text_area("List your education background (e.g., Degree, University, Dates)", height=150, value="", key="creator_education")

    st.markdown("### 💼 Work Experience") # Corrected emoji
    experience = st.text_area("List your work experience (e.g., Job Title, Company, Dates, Responsibilities/Achievements)", height=200, value="", key="creator_experience")

    st.markdown("### 🛠️ Skills") # Corrected emoji
    skills_input = st.text_area("List your key skills (separated by commas)", value="", key="creator_skills")

    if st.button("Generate Resume", key="creator_generate_button"):
        if not name or not (email or phone or linkedin):
             st.warning("Please fill in at least Name and some Contact Information.")
        else:
            from docx import Document
            doc = Document()

            # Add Name and Contact Info
            doc.add_heading(name, 0)
            contact_parts = []
            if email: contact_parts.append(email)
            if phone: contact_parts.append(phone)
            if linkedin: contact_parts.append(linkedin)
            if contact_parts:
                doc.add_paragraph(" | ".join(contact_parts))
            doc.add_paragraph() # Add a blank line

            # Add Summary
            if summary.strip():
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(summary)
                doc.add_paragraph()

            # Add Education
            if education.strip():
                doc.add_heading("Education", level=1)
                doc.add_paragraph(education)
                doc.add_paragraph()

            # Add Experience
            if experience.strip():
                doc.add_heading("Work Experience", level=1)
                doc.add_paragraph(experience)
                doc.add_paragraph()

            # Add Skills
            if skills_input.strip():
                doc.add_heading("Skills", level=1)
                doc.add_paragraph(skills_input)
                doc.add_paragraph()


            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📄 Download Resume", # Corrected emoji
                data=buffer,
                file_name=f"{name.replace(' ', '_').lower()}_resume.docx" if name else "my_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="creator_download_button"
            )
            st.success("Resume generated! Download the DOCX file.")


# === Resume-to-Job Matching Mode ===
elif app_mode == "Resume-to-Job Matching":

    col1, col2 = st.columns([2, 1])
    with col1:
        # Upload resume section
        st.markdown('<div class="sub-header">📄 Upload Your Resume</div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("", type=["pdf", "docx"], key="matching_uploader") # Added unique key

        # Process uploaded file and store in session state
        if uploaded_file:
            # Check if a new file was uploaded since the last process
            if uploaded_file.name != st.session_state.get('last_matching_upload_name'):
                 st.session_state.resume_text = "" # Clear previous data if new file
                 st.session_state.resume_structure_analysis = None
                 st.session_state.resume_skills_analysis = None
                 st.session_state.skills_extracted = []
                 st.session_state.job_results = None # Clear previous job results
                 # Update the last uploaded file name in session state
                 st.session_state.last_matching_upload_name = uploaded_file.name


            # Process the file if resume_text is empty (either first upload or new file)
            if not st.session_state.resume_text:
                with st.spinner("Processing resume..."):
                    resume_text = ""
                    if uploaded_file.name.endswith(".pdf"):
                        resume_text = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.name.endswith(".docx"):
                        resume_text = extract_text_from_docx(uploaded_file)
                    else:
                        st.error("Unsupported file format.")

                    # Store extracted text in session state
                    st.session_state.resume_text = resume_text

                    # Perform analysis only if text extraction was successful
                    if isinstance(st.session_state.resume_text, str) and not st.session_state.resume_text.startswith("Error extracting"):
                         # Use the analyze_resume function from resume_parser for structure
                         st.session_state.resume_structure_analysis = analyze_resume_structure(st.session_state.resume_text)
                         # Use the extract_skills function from semantic_matcher
                         st.session_state.skills_extracted = extract_skills_from_resume(st.session_state.resume_text)
                         # You might still want a separate skills analysis result if needed later
                         # st.session_state.resume_skills_analysis = analyze_resume_skills(st.session_state.resume_text) # If analyze_resume_skills still exists and is useful

                # Provide feedback after processing
                if isinstance(st.session_state.resume_text, str) and not st.session_state.resume_text.startswith("Error extracting") and st.session_state.resume_text.strip():
                    st.success("Resume processed successfully!")
                elif isinstance(st.session_state.resume_text, str) and st.session_state.resume_text.startswith("Error extracting"):
                     st.error(f"Failed to extract text: {st.session_state.resume_text}")
                else:
                    st.error("Could not extract text from the resume. Please check the file format.")


            # Display extracted text in an expandable section if available
            if st.session_state.resume_text and not st.session_state.resume_text.startswith("Error extracting"):
                 with st.expander("View Extracted Resume Text", expanded=False):
                     st.text_area("", st.session_state.resume_text, height=200, disabled=True)


    with col2:
        # Display resume stats if structure analysis was successful
        if st.session_state.resume_structure_analysis is not None and not st.session_state.resume_structure_analysis.get('error'):
            st.markdown('<div class="sub-header">📊 Resume Stats</div>', unsafe_allow_html=True)

            # Resume completeness score
            completeness = st.session_state.resume_structure_analysis.get('completeness_score', 0)
            st.markdown(f"**Resume Completeness:** {completeness}%")
            st.progress(completeness/100)

            # Resume sections found
            st.markdown("**Detected Sections:**")
            sections = st.session_state.resume_structure_analysis.get('sections', {})
            # Filter out 'other' and 'content' unless they have significant content
            detected_sections = [s for s in sections if s not in ['other', 'content'] or (s in ['other', 'content'] and len(sections[s].strip()) > 50)]
            if detected_sections:
                for section in detected_sections:
                    st.markdown(f"- {section.replace('_', ' ').title()}")
            else:
                st.info("No standard sections detected.")

            # Skills extracted
            st.markdown("**Extracted Skills:**")
            if st.session_state.skills_extracted:
                skills_text = ", ".join(st.session_state.skills_extracted[:15]) # Show more skills
                if len(st.session_state.skills_extracted) > 15:
                    skills_text += f" (+{len(st.session_state.skills_extracted) - 15} more)"
                st.markdown(skills_text)
            else:
                st.markdown("No specific skills detected. Consider adding a 'Skills' section.")
        elif uploaded_file and (st.session_state.resume_structure_analysis is None or st.session_state.resume_structure_analysis.get('error')):
             st.warning("Could not perform resume analysis. Check extracted text for errors.")
        elif not uploaded_file:
             st.info("Upload your resume to see analysis stats.")


    # Job matching section
    # Only show matching options if resume processing and job data loading were successful
    if st.session_state.resume_text and not st.session_state.resume_text.startswith("Error extracting") and st.session_state.resume_structure_analysis is not None and job_data_loaded:
        st.markdown('<div class="sub-header">🔎 Find Matching Jobs</div>', unsafe_allow_html=True)

        col1, col2 = st.columns([3, 1])

        with col1:
            # Job sector filter
            selected_sectors = st.multiselect(
                "Filter by job sector(s)",
                options=list(sector_options.keys()),
                placeholder="Choose one or more sectors...",
                key="matching_sector_filter" # Added key
            )

            # Convert sector selections to keywords
            selected_keywords = []
            if selected_sectors:
                for sector in selected_sectors:
                    selected_keywords.extend(sector_options[sector])

        with col2:
            # Number of results to show
            top_n = st.slider("Number of results", min_value=5, max_value=50, value=10, step=5, key="matching_top_n_slider") # Increased max, added key

            # Match button
            find_jobs = st.button("🔍 Find Matching Jobs", type="primary", use_container_width=True, key="find_jobs_button")


        if find_jobs:
            # Show progress and perform matching
            with st.spinner("Matching your profile to jobs..."):
                # Filter by selected keywords if any
                filtered_jobs = jobs_df.copy() # Work on a copy
                if selected_keywords:
                    # Create a case-insensitive regex pattern
                    # Ensure keywords are treated as whole words where appropriate, or use semantic filtering
                    # For simple keyword presence, a regex like this works, but might match parts of words
                    # pattern = '|'.join([re.escape(kw) for kw in selected_keywords]) # Use re.escape for special chars
                    # Using word boundaries \b might be better for whole words, but more complex with phrases
                    # Let's stick to the current pattern matching anywhere for broader filtering
                    pattern = '|'.join([re.escape(kw) for kw in selected_keywords])
                    filtered_jobs = filtered_jobs[
                        filtered_jobs["Job title"].str.contains(pattern, na=False, case=False) |
                        filtered_jobs["Job description"].str.contains(pattern, na=False, case=False) |
                        filtered_jobs["Requirements"].str.contains(pattern, na=False, case=False)
                    ]

                # If no jobs match the filters, show a message
                if filtered_jobs.empty:
                    st.warning("No jobs found matching your selected sectors. Try selecting different sectors or removing filters.")
                    st.session_state.job_results = pd.DataFrame() # Set to empty df
                else:
                    # Match using semantic embeddings
                    start_time = time.time()
                    # Ensure match_score is added by the semantic_match_resume function
                    # Pass the extracted resume text directly
                    results = semantic_match_resume(st.session_state.resume_text, filtered_jobs, top_n=top_n)
                    matching_time = time.time() - start_time

                    # Store results in session state
                    st.session_state.job_results = results

                    # Summary stats
                    if not results.empty:
                         st.success(f"Found {len(results)} matching jobs in {matching_time:.2f} seconds")
                    else:
                         st.info(f"No jobs matched your resume based on the semantic matching algorithm in {matching_time:.2f} seconds.")


                # Display visualization if results are available
                if st.session_state.job_results is not None and not st.session_state.job_results.empty:
                    # Create interactive visualizations of match scores
                    fig = px.bar(
                        st.session_state.job_results,
                        x='Job title',
                        y='match_score',
                        color='match_score',
                        color_continuous_scale='viridis',
                        labels={'match_score': 'Match Score (%)'},
                        title='Top Job Match Scores'
                    )
                    fig.update_layout(xaxis_tickangle=-45, height=400, margin=dict(t=40, b=150)) # Adjust layout
                    st.plotly_chart(fig, use_container_width=True)
                elif st.session_state.job_results is not None and st.session_state.job_results.empty and not filtered_jobs.empty:
                     st.info("The semantic matching algorithm did not find strong matches among the filtered jobs.")


        # Display job results if available in session state
        if st.session_state.job_results is not None and not st.session_state.job_results.empty:
            st.markdown('<div class="sub-header">✅ Top Matching Jobs</div>', unsafe_allow_html=True)

            # Get resume skills for matching
            # Use skills from session state analysis results for consistency
            resume_skills_list = st.session_state.skills_extracted # This is already a list from session state

            # Create tabs for different views
            tab1, tab2 = st.tabs(["List View", "Detailed View"])

            with tab1:
                # Display URL as clickable link in DataFrame - FIX APPLIED HERE
                def make_clickable_link(url):
                    if url and isinstance(url, str) and url.startswith('http'):
                        return f'<a href="{url}" target="_blank">{url}</a>'
                    return url # Return as is if not a valid http link or not a string

                # Apply the function to the URL column and display with unsafe_allow_html
                # Modify the DataFrame column directly BEFORE passing to st.dataframe
                filtered_jobs_display = st.session_state.job_results[['Job title', 'Company', 'Salary', 'match_score', 'URL']].copy()
                filtered_jobs_display['URL'] = filtered_jobs_display['URL'].apply(make_clickable_link)

                st.dataframe(
                    filtered_jobs_display,
                    use_container_width=True
                )

            with tab2:
                # Show detailed job cards
                # Sort results by match score descending for detailed view
                detailed_results = st.session_state.job_results.sort_values(by='match_score', ascending=False)

                for i, (_, row) in enumerate(detailed_results.iterrows()):
                    # Combine job description and requirements for skill matching
                    job_text = str(row.get("Job description", "")) + " " + str(row.get("Requirements", ""))

                    # Get matched and missing skills if resume skills are available
                    matched_skills, missing_skills = [], []
                    # Ensure resume_skills_list is not empty before calling get_skill_matches
                    if resume_skills_list:
                        matched_skills, missing_skills = get_skill_matches(resume_skills_list, job_text)

                    # Format match score with color
                    match_score = row['match_score']
                    if match_score >= 80:
                        match_class = "match-score-high"
                    elif match_score >= 60:
                        match_class = "match-score-medium"
                    else:
                        match_class = "match-score-low"

                    # Job card
                    st.markdown(f'<div class="match-section">', unsafe_allow_html=True)
                    st.markdown(f'<div class="job-title">{i+1}. {row["Job title"]}</div>', unsafe_allow_html=True)
                    st.markdown(f"**Company:** {row.get('Company', 'Unknown')}")
                    st.markdown(f"**Salary:** {row.get('Salary', 'Not specified')}")
                    st.markdown(f'**Match Score:** <span class="{match_class}">{match_score:.1f}%</span>', unsafe_allow_html=True)

                    # Job URL - Check if URL is valid before displaying
                    job_url = row.get('URL', '#')
                    if job_url and isinstance(job_url, str) and job_url != '#' and job_url.startswith('http'): # Basic check and type check
                         st.markdown(f"[🔗 View Job Posting]({job_url})")
                    elif job_url and isinstance(job_url, str) and job_url != '#':
                         st.markdown(f"URL: {job_url}") # Display if it exists but isn't a standard http link
                    else:
                         st.markdown("URL: Not available")


                    # Matched and Missing Skills section
                    # Only show this section if skills were extracted from resume AND job text is available
                    if resume_skills_list and job_text.strip():
                        if matched_skills:
                            st.markdown('<div class="matched-keywords">🟢 **Matched skills:** ' +
                                      ', '.join(matched_skills[:10]) + # Show more skills
                                      (f' (+{len(matched_skills)-10} more)' if len(matched_skills) > 10 else '') +
                                      '</div>', unsafe_allow_html=True)

                        # Only show missing skills if there are any
                        if missing_skills:
                            st.markdown('<div class="missing-keywords">🔴 **Potentially missing skills:** ' + # Improved wording
                                      ', '.join(missing_skills[:10]) + # Show more skills
                                      (f' (+{len(missing_skills)-10} more)' if len(missing_skills) > 10 else '') +
                                      '</div>', unsafe_allow_html=True)
                        # Add a message if no missing skills were found but skills were extracted
                        elif matched_skills: # If there were matched skills, but no missing ones
                             st.markdown('<div class="matched-keywords">✅ Your skills seem to cover key requirements for this job!</div>', unsafe_allow_html=True)
                        else: # If resume skills were extracted, but neither matched nor missing were found
                             st.info("Could not identify specific skill matches or gaps for this job.")

                    elif resume_skills_list and not job_text.strip():
                         st.info("Job description/requirements are empty, cannot analyze skill match.")
                    elif not resume_skills_list:
                         st.info("Upload your resume to see skill matches for jobs.")


                    # Expandable job description
                    with st.expander("View Job Details"):
                        st.markdown("#### Job Description")
                        st.markdown(row.get("Job description", "No description available"))

                        st.markdown("#### Requirements")
                        st.markdown(row.get("Requirements", "No requirements specified"))

                    st.markdown('</div>', unsafe_allow_html=True)
                    st.markdown("---")

            # Download results
            output = io.BytesIO()
            # Ensure 'match_score' column is included in the download
            cols_to_download = ['Job title', 'Company', 'Salary', 'match_score', 'URL', 'Job description', 'Requirements']
            # Filter results dataframe to only include these columns
            download_df = st.session_state.job_results.copy()
            # Ensure all columns exist in the download_df before selecting
            download_cols_present = [col for col in cols_to_download if col in download_df.columns]
            download_df = download_df[download_cols_present]


            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                download_df.to_excel(writer, index=False)
            output.seek(0)

            st.download_button(
                label="⬇ Download results as Excel",
                data=output,
                file_name="matched_jobs.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.document",
                key="matching_download_button"
            )

            # Resume improvement tips based on job matches
            st.markdown('<div class="sub-header">🚀 Improve Your Resume for Better Matches</div>', unsafe_allow_html=True)
            st.markdown('<div class="tips-section">', unsafe_allow_html=True)
            st.markdown("Based on the skills found in the top matching jobs but potentially missing from your resume, here are areas to consider adding or highlighting:")

            # Generate dynamic tips based on job results and missing skills
            all_missing_skills = []
            # Only analyze missing skills if resume skills were extracted and job results exist
            if resume_skills_list and st.session_state.job_results is not None and not st.session_state.job_results.empty:
                for _, row in st.session_state.job_results.iterrows():
                    job_text = str(row.get("Job description", "")) + " " + str(row.get("Requirements", ""))
                    # Ensure job_text is not empty before calling get_skill_matches
                    if job_text.strip():
                         _, missing = get_skill_matches(resume_skills_list, job_text)
                         all_missing_skills.extend(missing)

            # Count most common missing skills (limit to top 10-15)
            common_missing = Counter(all_missing_skills).most_common(15) # Show more common skills

            if common_missing:
                st.markdown("#### Top Skills to Consider Adding or Highlighting:")
                st.info("These skills appeared frequently in the job descriptions and requirements of the filtered jobs, suggesting common skills or qualifications employers are seeking.")
                cols = st.columns(3) # Use columns for listing skills
                for i, (skill, count) in enumerate(common_missing):
                     cols[i % 3].markdown(f"- **{skill}** (mentioned in {count} job{'s' if count > 1 else ''})")

            else:
                 st.info("Based on the top matches, no specific skills were consistently missing from your resume's extracted skill list. Ensure your existing skills are clearly articulated.")


            # General resume improvement tips (can link to Resume Analysis tab)
            st.markdown("#### General Tips for Job Matching:")
            st.markdown("""
            - **Include a Skills Section:** A clear section listing your technical and soft skills helps matching algorithms.
            - **Use Keywords:** Incorporate keywords found in job descriptions throughout your resume, especially in your experience bullet points.
            - **Tailor for Each Job:** Adapt your resume slightly for each application, emphasizing the skills and experiences most relevant to that specific role.
            - **Quantify:** Use numbers to describe your achievements in previous roles, making your impact clear.
            """)
            st.markdown('</div>', unsafe_allow_html=True)

        elif st.session_state.job_results is not None and st.session_state.job_results.empty and find_jobs:
             st.info("No jobs were found matching your criteria after filtering or during the matching process.")

    elif uploaded_file and (isinstance(st.session_state.resume_text, str) and st.session_state.resume_text.startswith("Error extracting") or st.session_state.resume_structure_analysis is None):
        st.warning("Please upload a valid resume file (PDF or DOCX) and ensure text extraction was successful to find matching jobs.")
    elif not uploaded_file:
         st.info("Upload your resume above to find matching jobs based on your profile.")
    elif not job_data_loaded:
         st.warning("Job data could not be loaded. Job matching is unavailable.")


# === Resume Analysis Mode (Enhanced) ===
elif app_mode == "Resume Analysis":

    st.markdown('📊 Resume Analyzer', unsafe_allow_html=True)
    st.write("Upload your resume to get a detailed analysis of its content and effectiveness.")

    # Initialize analysis results to None before the file upload
    # These are local variables for this mode, they will be populated if a file is uploaded and processed
    resume_structure_analysis = None
    skills_analysis = None
    resume_text_analysis = "" # Local variable for text in this mode

    # Upload resume for analysis
    uploaded_file = st.file_uploader("", type=["pdf", "docx"], key="analysis_uploader_mode") # Added unique key


    if uploaded_file:
        # Check if a new file was uploaded since the last process in this mode
        if uploaded_file.name != st.session_state.get('last_analysis_upload_name'):
             # Clear previous local variables if a new file
             resume_structure_analysis = None
             skills_analysis = None
             resume_text_analysis = ""
             # Update the last uploaded file name in session state for this mode
             st.session_state.last_analysis_upload_name = uploaded_file.name


        # Process the file if resume_text_analysis is empty (either first upload or new file)
        if not resume_text_analysis: # Check the local variable
             with st.spinner("Analyzing your resume..."):
                extracted_text = ""
                if uploaded_file.name.endswith(".pdf"):
                    extracted_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.name.endswith(".docx"):
                    extracted_text = extract_text_from_docx(uploaded_file)
                else:
                    st.error("Unsupported file format.")

                # Store extracted text in the local variable
                resume_text_analysis = extracted_text

                # Analyze resume only if text was extracted and is not an error message
                if isinstance(resume_text_analysis, str) and not resume_text_analysis.startswith("Error extracting") and resume_text_analysis.strip():
                    # Use the analyze_resume function from resume_parser for structure
                    resume_structure_analysis = analyze_resume_structure(resume_text_analysis)
                    # Use the extract_skills and extract_keywords from semantic_matcher
                    skills_analysis = {
                        "skills": extract_skills_from_resume(resume_text_analysis),
                        "keywords": extract_resume_keywords(resume_text_analysis)
                    }
                    st.success("Resume analyzed successfully!")
                elif isinstance(resume_text_analysis, str) and resume_text_analysis.startswith("Error extracting"):
                     st.error(f"Failed to extract text: {resume_text_analysis}")
                else:
                    st.error("Could not extract text from the resume for analysis.")

        # If text was already extracted (e.g., on rerun), use the existing local variables
        # This assumes analysis results are implicitly available if text_analysis is not empty
        # A more robust way would be to store analysis results in session state for this mode too
        # But for simplicity, we'll rely on the processing block above running on file upload/change.
        # If the app reruns for other reasons without a new file, analysis won't re-run unless needed.
        # Let's re-run analysis if text is present but analysis results are not (e.g., after a code change)
        if isinstance(resume_text_analysis, str) and not resume_text_analysis.startswith("Error extracting") and resume_text_analysis.strip() and (resume_structure_analysis is None or skills_analysis is None):
             with st.spinner("Re-analyzing resume..."):
                 resume_structure_analysis = analyze_resume_structure(resume_text_analysis)
                 skills_analysis = {
                     "skills": extract_skills_from_resume(resume_text_analysis),
                     "keywords": extract_resume_keywords(resume_text_analysis)
                 }
                 st.success("Resume re-analyzed successfully!")


    # Display analysis results only if analysis was successful and results are not error messages
    if resume_structure_analysis is not None and skills_analysis is not None and not resume_structure_analysis.get('error'):
        # Create tabs for different analysis views
        tab1, tab2, tab3 = st.tabs(["Overview", "Content Details", "Improvement Suggestions"]) # Renamed tabs

        with tab1:
            st.markdown("#### Summary Overview")

            # Overview metrics
            col1, col2, col3 = st.columns(3)

            # Completeness score
            completeness = resume_structure_analysis.get('completeness_score', 0)
            with col1:
                st.metric("Completeness Score", f"{completeness}%")
                st.progress(completeness/100)
                if completeness < 50:
                    st.warning("Your resume seems very incomplete. Consider adding key sections.")
                elif completeness < 80:
                     st.info("Your resume could be more complete. Ensure all relevant sections are present.")


            # Section count
            sections = resume_structure_analysis.get('sections', {})
            # Filter out 'other' and 'content' which are usually catch-alls unless they have significant content
            detected_sections = [s for s in sections if s not in ['other', 'content'] or (s in ['other', 'content'] and len(sections[s].strip()) > 50)]
            section_count = len(detected_sections)
            with col2:
                 st.metric("Key Sections Found", section_count)
                 if section_count < 3: # Basic check for core sections like Contact, Education, Experience
                     st.warning("Fewer key sections detected than expected. Ensure you have Contact, Education, and Experience sections.")


            # Skills count
            skills = skills_analysis.get('skills', [])
            skills_count = len(skills)
            with col3:
                 st.metric("Skills Detected", skills_count)
                 if skills_count < 10:
                     st.info("Detecting more specific skills can improve job matching.")

            st.markdown("---") # Separator

            # Contact information
            st.markdown("#### Contact Information")
            contact_info = resume_structure_analysis.get('contact_info', {})

            if contact_info.get('emails') or contact_info.get('phones') or contact_info.get('linkedin') or contact_info.get('locations') or contact_info.get('name'):
                # Display Name if found
                if contact_info.get('name'):
                     st.write(f"👤 **Name:** {contact_info['name'][0]}")
                if contact_info.get('emails'):
                    st.write(f"📧 **Email:** {contact_info['emails'][0]}")
                if contact_info.get('phones'):
                    st.write(f"📱 **Phone:** {contact_info['phones'][0]}")
                if contact_info.get('linkedin'):
                    st.write(f"🔗 **LinkedIn:** {contact_info['linkedin'][0]}")
                if contact_info.get('locations'):
                     st.write(f"📍 **Location:** {', '.join(contact_info['locations'])}")
            else:
                st.error("❌ **No or incomplete contact information detected.** Ensure your resume includes clear Name, Email, Phone, and ideally LinkedIn/Location.")


        with tab2: # Renamed to Content Details
            st.markdown("#### Content Breakdown")

            # Skills display
            st.markdown("##### 🛠️ Detected Skills")
            if skills:
                st.info("These are the specific skills identified in your resume. Ensure they are relevant and clearly stated.")
                # Use columns for a more organized tag-like display
                cols = st.columns(5) # Use more columns
                for i, skill in enumerate(skills):
                    cols[i % 5].markdown(f"- `{skill}`") # Using markdown code formatting for clarity
            else:
                st.warning("⚠️ **No specific skills detected.** Add a dedicated 'Skills' section or sprinkle keywords throughout your experience.")

            st.markdown("---") # Separator

            # Keywords visualization (Word Cloud)
            st.markdown("##### 🔑 Key Terms and Phrases")
            st.info("This word cloud shows the most frequent and potentially relevant terms found in your resume, giving you an idea of its focus.")

            keywords = skills_analysis.get('keywords', [])
            if keywords:
                # Join keywords into a single string for the word cloud
                text_for_wordcloud = " ".join(keywords)

                # Create and display word cloud
                try:
                    # Ensure WordCloud and matplotlib are imported at the top
                    wordcloud = WordCloud(width=800, height=400, background_color='white',
                                        colormap='viridis', max_words=50, min_font_size=10).generate(text_for_wordcloud)

                    fig, ax = plt.subplots(figsize=(10, 5))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis('off')
                    st.pyplot(fig)

                except Exception as e: # Catch any exception during word cloud generation
                    st.error(f"Error generating word cloud: {e}")

            else:
                st.info("No significant keywords detected for visualization.")

            st.markdown("---") # Separator

            # Display resume sections
            st.markdown("##### 📄 Resume Sections Content")
            st.info("Review the extracted content for each section to ensure accuracy and completeness.")
            # Filter out 'other' and 'content' from sections display unless they have substantial text
            sections_to_display = {k: v for k, v in sections.items() if k not in ['other', 'content'] or (k in ['other', 'content'] and len(sections[k].strip()) > 50)}

            if sections_to_display:
                for section_name, content in sections_to_display.items():
                    # Use a slightly better title for expanders
                    section_display_name = section_name.replace('_', ' ').title()
                    with st.expander(f"{section_display_name} Section"):
                        if content.strip():
                             st.text(content.strip()) # Use st.text to preserve formatting
                        else:
                             st.info(f"This section was detected but appears empty.")

            else:
                st.warning("⚠️ **No clear, standard sections detected in your resume.** Using standard headings (like 'Education', 'Work Experience', 'Skills') is crucial for ATS.")

        with tab3: # Renamed to Improvement Suggestions
            st.markdown("#### Improve Your Resume")
            st.info("Here are tailored suggestions to enhance your resume based on the analysis.")

            # --- Tailored Suggestions based on Analysis ---

            st.markdown("##### Based on Analysis:")

            # Completeness suggestions
            if completeness < 70:
                st.markdown("###### 🔴 Completeness & Structure:")
                st.markdown(f"- Your resume completeness score is **{completeness}%**. Key sections might be missing or brief.")
                st.markdown("- Ensure you have dedicated sections for: **Contact Information, Summary/Objective, Work Experience, Education, and Skills.**")
                # Check for specific missing core sections
                core_sections_keys = ['contact_info', 'summary', 'education', 'experience', 'skills'] # Added summary to core check
                missing_core_sections = []
                for section_key in core_sections_keys:
                     if section_key == 'contact_info':
                          # Check if contact_info dictionary is empty or all its relevant lists are empty
                          if not any(resume_structure_analysis.get('contact_info', {}).values()):
                               missing_core_sections.append("Contact Information")
                     # Check if the section key is NOT in the detected sections OR if its content is too short
                     elif section_key not in sections or len(sections.get(section_key, '').strip()) < 50: # Basic length check
                          missing_core_sections.append(section_key.replace('_', ' ').title())

                if missing_core_sections:
                     st.markdown(f"- Specifically, consider adding or expanding sections like: **{', '.join(missing_core_sections)}**.")
                else:
                     st.info("Core sections (Contact, Summary, Education, Experience, Skills) seem to be present.")


            # Skills suggestions
            skills = skills_analysis.get('skills', []) # Get skills again from the analysis result
            skills_count = len(skills)
            if skills_count < 15: # Slightly higher threshold for actionable advice
                 st.markdown("###### 🔶 Skills Visibility:")
                 st.markdown(f"- Only **{skills_count}** distinct skills were easily detected.")
                 st.markdown("- Explicitly list your technical skills, software proficiency, and relevant soft skills in a dedicated 'Skills' section.")
                 st.markdown("- Review job descriptions for roles you want and incorporate relevant keywords into your skills and experience sections.")
            else:
                 st.info("A good number of skills were detected in your resume.")


            # Content depth check (basic based on total text length)
            if isinstance(resume_text_analysis, str) and len(resume_text_analysis) < 1000 and completeness >= 50: # Avoid this tip if completeness is already very low
                 st.markdown("###### 📝 Content Depth:")
                 st.markdown("- Your resume text is relatively brief. Expand on your accomplishments in your experience section.")
                 st.markdown("- Use detailed bullet points that describe **what you did**, **how you did it**, and **the positive result (quantify!)**.")
            elif isinstance(resume_text_analysis, str) and len(resume_text_analysis) >= 1000:
                 st.info("Your resume seems to have a good level of detail.")


            st.markdown("---") # Separator

            # --- General Best Practices & ATS Tips ---

            st.markdown("##### General Best Practices & ATS Tips:")
            st.info("These tips help improve your resume's readability for both recruiters and automated systems (ATS).")

            st.markdown("""
            - **Quantify Achievements:** Whenever possible, use numbers, percentages, or data points to describe your impact (e.g., "Increased efficiency by 20%", "Managed a team of 5").
            - **Use Action Verbs:** Start bullet points with strong action verbs (e.g., *Led, Developed, Managed, Created, Implemented, Analyzed*).
            - **Tailor Your Resume:** Modify your resume slightly for each job application by incorporating keywords from the job description.
            - **ATS Formatting:**
                - Use standard resume section titles (Education, Work Experience, Skills, Projects, etc.). Avoid creative or unusual headings.
                - Avoid complex designs: no tables, multi-column layouts (especially in the main body), headers/footers (ATS might ignore them), text boxes, or images.
                - Stick to common, readable fonts (Arial, Calibri, Times New Roman, Verdana) and a font size between 10-12pt.
                - Use reverse chronological order for experience and education.
                - Save as a `.docx` or `.pdf` (check application instructions; PDF is generally safe if not image-based).
            - **Proofread:** Thoroughly check for typos, grammatical errors, and inconsistent formatting. A single error can make a negative impression.
            """)

            st.markdown("---") # Separator

            st.markdown("##### 📚 Additional Resources:")
            st.markdown("""
            - Review successful resume examples in your target industry.
            - Use tools like Grammarly to check for writing errors.
            - Consider getting feedback from peers or career services.
            """)

    elif uploaded_file and (isinstance(resume_text_analysis, str) and resume_text_analysis.startswith("Error extracting")):
        st.error("Could not extract readable text from the uploaded file. Please ensure it's a searchable PDF or a standard DOCX.")
    elif uploaded_file and (resume_structure_analysis is None or skills_analysis is None):
         st.warning("Could not complete resume analysis. Please check the extracted text.")
    elif not uploaded_file:
         st.info("Upload your resume above to get a detailed analysis.")


# === Job Market Explorer Mode ===
elif app_mode == "Job Market Explorer":

    st.markdown('🔍 Job Market Explorer', unsafe_allow_html=True)
    st.write("Explore the current job market trends and opportunities based on the loaded data.")

    if not job_data_loaded:
         # Message already shown above
         pass
    else:
        # Start with the full dataset
        filtered_jobs = jobs_df.copy()

        # Apply company filter
        if 'selected_companies' in locals() and selected_companies:
            filtered_jobs = filtered_jobs[filtered_jobs['Company'].isin(selected_companies)]

        # Apply salary filter
        if 'salary_min' in locals() and salary_min > 0:
            # Ensure 'salary_value' column is created for filtering
            def extract_salary_value(salary_str):
                if pd.isna(salary_str) or str(salary_str).strip().lower() == 'not specified':
                    return 0
                # Extract numbers, handling potential commas
                numbers = re.findall(r'\d+', str(salary_str).replace(',', ''))
                if numbers:
                    # Take the first number found as the minimum value
                    return int(numbers[0])
                return 0

            # Apply extraction and filter
            filtered_jobs['salary_value'] = filtered_jobs['Salary'].apply(extract_salary_value)
            filtered_jobs = filtered_jobs[filtered_jobs['salary_value'] >= salary_min]
            filtered_jobs = filtered_jobs.drop(columns=['salary_value']) # Drop the helper column

        # Apply sector filter
        if 'selected_job_sectors' in locals() and selected_job_sectors:
            # Convert selected sectors back to keywords
            selected_explorer_keywords = []
            for sector in selected_job_sectors:
                 selected_explorer_keywords.extend(sector_options[sector])

            if selected_explorer_keywords:
                # Create a case-insensitive regex pattern from selected keywords
                pattern = '|'.join([re.escape(kw) for kw in selected_explorer_keywords])
                filtered_jobs = filtered_jobs[
                    filtered_jobs["Job title"].str.contains(pattern, na=False, case=False) |
                    filtered_jobs["Job description"].str.contains(pattern, na=False, case=False) |
                    filtered_jobs["Requirements"].str.contains(pattern, na=False, case=False)
                ]


        # Show statistics and visualizations
        if filtered_jobs.empty:
            st.warning("No jobs found matching the current filters. Try adjusting your filters in the sidebar.")
        else:
            st.markdown(f"**Displaying {len(filtered_jobs)} jobs matching your filters.**")
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### Job Count by Sector")

                # Classify jobs into sectors for visualization *of the filtered data*
                @st.cache_data(ttl=300) # Cache classification for performance if filters change frequently
                def classify_jobs_for_viz(df_subset, sector_keywords):
                     # Helper function for classification
                     def classify_job_single(job_title, job_text, sector_keywords):
                         text = (str(job_title) + " " + str(job_text)).lower()
                         for sector, keywords in sector_keywords.items():
                             # Check if any keyword from the sector is in the text
                             if any(keyword.lower() in text for keyword in keywords):
                                 return sector
                         return "Other"

                     df_subset['sector'] = df_subset.apply(
                         lambda row: classify_job_single(row['Job title'], str(row.get('Job description', '')) + " " + str(row.get('Requirements', '')) , sector_keywords), # Combine desc/req
                         axis=1
                     )
                     return df_subset['sector'].value_counts()

                sector_counts = classify_jobs_for_viz(filtered_jobs.copy(), sector_options)


                # Create pie chart
                if not sector_counts.empty:
                    fig = px.pie(
                        names=sector_counts.index,
                        values=sector_counts.values,
                        title="Job Distribution by Sector",
                        hole=0.4,
                        color_discrete_sequence=px.colors.qualitative.Pastel
                    )
                    fig.update_layout(margin=dict(t=40, b=0, l=0, r=0)) # Adjust margins
                    st.plotly_chart(fig, use_container_width=True)
                else:
                     st.info("Not enough data to show sector distribution.")


            with col2:
                st.markdown("#### Top Companies Hiring")

                # Count jobs by company for the filtered data
                company_counts = filtered_jobs['Company'].value_counts().head(10)

                # Create horizontal bar chart
                if not company_counts.empty:
                    fig = px.bar(
                        x=company_counts.values,
                        y=company_counts.index,
                        orientation='h',
                        title="Top 10 Companies with Open Positions",
                        labels={'x': 'Number of Jobs', 'y': 'Company'},
                        color=company_counts.values,
                        color_continuous_scale='Viridis'
                    )
                    fig.update_layout(yaxis={'categoryarray': company_counts.index.tolist()}, margin=dict(t=40, b=0, l=0, r=0)) # Ensure correct order and margins
                    st.plotly_chart(fig, use_container_width=True)
                else:
                     st.info("Not enough data to show top companies.")


            # Show job listing table
            st.markdown("#### Available Job Listings")
            MAX_DISPLAY = 10  # Limit how many jobs are shown
            shown_jobs = filtered_jobs[['Job title', 'Company', 'Salary', 'URL']].head(MAX_DISPLAY)
            
            with st.expander(f"📄 Click to View Top {MAX_DISPLAY} Job Listings"):
                for _, row in shown_jobs.iterrows():
                    title = row.get('Job title', 'Unknown Job')
                    company = row.get('Company', 'Unknown Company')
                    salary = row.get('Salary', 'Not specified')
                    url = row.get('URL', '')

                    if url and isinstance(url, str) and url.startswith("http"):
                        st.markdown(f"- **{title}** at **{company}** – 💰 {salary} – [🔗 View Job Posting]({url})")
                    else:
                        st.markdown(f"- **{title}** at **{company}** – 💰 {salary} – No link available")

            if len(filtered_jobs) > MAX_DISPLAY:
                st.info(f"Only showing the first {MAX_DISPLAY} jobs. Refine filters to see more.")


            # Word cloud of job requirements keywords
            st.markdown("#### Most In-Demand Skills/Keywords")
            st.info("This word cloud highlights recurring terms in the job descriptions and requirements of the filtered jobs, suggesting common skills or qualifications employers are seeking.")

            # Extract and combine text from job descriptions and requirements
            all_job_text = " ".join(filtered_jobs['Job description'].fillna('') + " " + filtered_jobs['Requirements'].fillna(''))

            if all_job_text.strip():
                # Create word cloud
                try:
                    # Ensure WordCloud and matplotlib are imported at the top
                    wordcloud = WordCloud(width=800, height=400, background_color='white',
                                        colormap='plasma', max_words=100, min_font_size=10,
                                        stopwords=None, # You might want to add custom stopwords
                                        collocations=False).generate(all_job_text) # Set collocations=False to avoid combining common pairs

                    # Display the word cloud
                    fig, ax = plt.subplots(figsize=(10, 5))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis('off')
                    st.pyplot(fig)

                except Exception as e: # Catch any exception during word cloud generation
                     st.error(f"Error generating word cloud: {e}")
            else:
                st.info("No sufficient job description/requirements text available to generate a word cloud.")


# Footer
st.markdown("---")
st.write("Developed with ❤️ using Streamlit")
st.write("Data source: Zangia (Filtered public listings)")
