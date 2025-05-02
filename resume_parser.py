import fitz # PyMuPDF
import docx2txt
import tempfile
import re
import os
from docx import Document
import pandas as pd
import spacy
from collections import defaultdict

# Load spaCy model for NER
try:
    nlp = spacy.load('en_core_web_sm')
except:
    # Fallback for environments where direct download might be restricted
    # In a Streamlit environment, ensure this package is in your requirements.txt
    # or pre-installed. This subprocess call might not work everywhere.
    try:
        import subprocess
        subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"], check=True)
        nlp = spacy.load('en_core_web_sm')
        print("spaCy model downloaded and loaded successfully.")
    except Exception as e:
        print(f"Error loading or downloading spaCy model: {e}")
        print("Proceeding without spaCy NER features.")
        nlp = None # Set nlp to None if loading fails

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF with improved formatting preservation."""
    text = ""
    try:
        # Move file pointer to the beginning in case it was read before
        pdf_file.seek(0)
        with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
            for page in doc:
                # Use 'text' with 'textenabled=True' for potentially better block order
                # Or stick to 'blocks' but process them carefully
                # Let's refine the 'blocks' processing slightly
                blocks = page.get_text("blocks")
                # Sort blocks by their top-left corner (y then x) to handle columns
                blocks.sort(key=lambda block: (block[1], block[0]))
                for block in blocks:
                     block_text = block[4].strip()
                     if block_text:
                          text += block_text + "\n\n" # Add extra newline between blocks
    except Exception as e:
        # Log the error or handle it appropriately
        print(f"Error extracting PDF: {str(e)}")
        return f"Error extracting PDF: {str(e)}" # Return error message to caller
    return clean_resume_text(text)

def extract_text_from_docx(docx_file):
    """Extract text from DOCX with improved structure preservation."""
    try:
        # Move file pointer to the beginning
        docx_file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_file.read())
            tmp_path = tmp.name

        doc = Document(tmp_path)
        full_text = []

        # Extract header text
        for section in doc.sections:
             for header in section.header.paragraphs:
                  if header.text.strip():
                       full_text.append(header.text.strip())

        # Extract main body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text)) # Join cell text with a separator

        os.unlink(tmp_path) # Clean up the temporary file
        return clean_resume_text("\n\n".join(full_text)) # Join parts with double newline

    except Exception as e:
        # Log the error or handle it appropriately
        print(f"Error extracting DOCX: {str(e)}")
        return f"Error extracting DOCX: {str(e)}" # Return error message to caller


def clean_resume_text(text):
    """Basic text cleaning for resume content."""
    # Remove excessive whitespace, keeping some structure
    text = re.sub(r'[ \t]+', ' ', text) # Replace multiple spaces/tabs with a single space
    text = re.sub(r'\n\s*\n', '\n\n', text) # Replace multiple newlines with double newline
    text = text.replace('•', '\n• ') # Ensure bullet points are on new lines
    # Add a space between a lowercase letter and a following uppercase letter (e.g., "skillSet" -> "skill Set")
    # This can sometimes help split concatenated words, but be cautious as it might split valid words (e.g., "iPhone")
    # Let's make this optional or more targeted if needed. For now, keep it as it was.
    # text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text) # Keep this as per original code
    text = re.sub(r'\n{3,}', '\n\n', text) # Reduce excessive newlines

    return text.strip()


def extract_resume_sections(text):
    """
    Extracts content into standard resume sections based on common headers.
    Improved logic to collect content until the next header.
    """
    # More robust patterns including common variations and allowing for case-insensitivity
    section_patterns = {
        'contact': r'^\s*(contact|personal|info|information|profile|details)\s*$',
        'summary': r'^\s*(summary|objective|about)\s*$', # Added Summary/Objective
        'education': r'^\s*(education|academic|qualification|degree|university|school)\s*$',
        'experience': r'^\s*(experience|employment|work|history|professional|career)\s*$',
        'skills': r'^\s*(skills|abilities|expertise|competencies|proficiencies|technical skills|soft skills)\s*$',
        'projects': r'^\s*(projects|portfolio|works)\s*$',
        'achievements': r'^\s*(achievements|accomplishments|honors|awards)\s*$',
        'languages': r'^\s*(languages|linguistic)\s*$',
        'references': r'^\s*(references|recommendations)' # Removed $ anchor to allow trailing text on the header line
    }

    sections = defaultdict(str)
    lines = text.split('\n')
    current_section = 'other' # Default section for content before the first recognized header
    section_found = False # Flag to track if any standard section header was found

    # First pass: Identify section start lines and map them
    section_starts = {}
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        # Check if the line matches any section pattern and is reasonably short
        for section_name, pattern in section_patterns.items():
            # Use re.search instead of re.match to find pattern anywhere on the line
            # Keep word count check to avoid matching long sentences
            if re.search(pattern, line_lower, re.IGNORECASE) and len(line_lower.split()) < 5:
                section_starts[i] = section_name # Map line index to section name
                section_found = True
                break # Move to the next line once a header is matched

    # Get sorted list of line indices where sections start
    sorted_start_lines = sorted(section_starts.keys())

    # Second pass: Iterate through lines and assign content to sections
    current_section = 'other'
    section_start_index = 0 # Index for iterating through sorted_start_lines

    for i, line in enumerate(lines):
        # Check if the current line is the start of the next section in our sorted list
        if section_start_index < len(sorted_start_lines) and i == sorted_start_lines[section_start_index]:
            # Update the current section name
            current_section = section_starts[i]
            section_start_index += 1 # Move to the next expected section start line
            # Optionally, you might want to add the header line itself to the section content
            # sections[current_section] += line.strip() + "\n" # Decide if headers should be included
        else:
            # If it's not a header line, add it to the current section's content
            sections[current_section] += line.strip() + "\n"

    # Clean up whitespace within sections
    for section_name in sections:
        sections[section_name] = re.sub(r'\n{2,}', '\n\n', sections[section_name].strip()) # Reduce multiple newlines within sections

    # Handle case where no standard headers were found - put everything in 'content'
    # Only do this if 'other' is the only section found after processing
    if not section_found or (len(sections) == 1 and 'other' in sections and sections['other'].strip()):
         sections = {'content': text.strip()}
    else:
        # If 'other' section exists but contains only whitespace after cleaning, remove it
        if 'other' in sections and not sections['other'].strip():
             del sections['other']
        # If 'other' section exists and has content, keep it for unclassified text at the start/end
        pass


    return dict(sections) # Convert defaultdict back to dict

def extract_contact_info(text):
    """Extracts contact information using regex and spaCy NER."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    # More robust phone pattern, allowing various formats
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b|\b\d{7,}\b'
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    linkedin = re.findall(linkedin_pattern, text)

    locations = []
    name = []
    # Use spaCy only if the model loaded successfully
    if nlp:
        doc = nlp(text[:2000]) # Process a larger chunk for location/name
        locations = list(set([ent.text for ent in doc.ents if ent.label_ == "GPE"])) # Use set to get unique locations

        # Basic attempt to get a name from the first few lines if no contact section found
        # Also look for PERSON entities anywhere if a contact section IS found, but prioritize start
        doc_full = nlp(text[:500]) # Look at the first 500 characters for a name
        names_found = [ent.text for ent in doc_full.ents if ent.label_ == "PERSON"]
        if names_found:
             # Take the longest name found as a potential candidate
             name = [max(names_found, key=len)]


    return {
        'emails': emails,
        'phones': phones,
        'linkedin': linkedin,
        'locations': locations,
        'name': name # Include name extraction here
    }

def analyze_resume(text):
    """Analyzes resume text for structure, contact info, and completeness."""
    # Check if text extraction failed
    if isinstance(text, str) and text.startswith("Error extracting"):
        return {'error': text}
    if not isinstance(text, str) or not text.strip():
         return {'error': "No text extracted from resume."}


    sections = extract_resume_sections(text)
    contact_info = extract_contact_info(text)

    # Use spaCy only if the model loaded successfully
    entities = {}
    if nlp:
        doc = nlp(text[:5000]) # Process a larger chunk for entities
        entities = {
            'organizations': list(set([ent.text for ent in doc.ents if ent.label_ == "ORG"])),
            'dates': list(set([ent.text for ent in doc.ents if ent.label_ == "DATE"])),
            'people': list(set([ent.text for ent in doc.ents if ent.label_ == "PERSON"])) # Includes potential name
        }

    # Calculate completeness score based on presence and length of key sections
    # Adjusted weights slightly to make 100 achievable with core sections + contact
    section_weights = {
        'contact': 15, # Increased weight for contact info presence
        'summary': 10,
        'education': 20,
        'experience': 30,
        'skills': 20,
        'projects': 5,
        'achievements': 5,
        'languages': 5
    }

    completeness_score = 0

    # Add score for contact info presence
    # Max contact score = 15 (5 for email, 5 for phone, 3 for linkedin, 2 for location)
    if contact_info.get('emails'): completeness_score += 5
    if contact_info.get('phones'): completeness_score += 5
    if contact_info.get('linkedin'): completeness_score += 3
    if contact_info.get('locations'): completeness_score += 2
    if contact_info.get('name'): completeness_score += 3 # Add score for name detection


    # Add score for sections based on presence and minimum content length
    min_section_length = 50 # Minimum characters to consider a section present and meaningful
    for section, weight in section_weights.items():
        # Check if the section was detected AND has sufficient content
        # Exclude 'contact' from this length check as its presence is scored above
        if section != 'contact' and section in sections and len(sections[section].strip()) > min_section_length:
             completeness_score += weight

    # Cap score at 100
    completeness_score = min(completeness_score, 100)


    return {
        'sections': sections, # Dictionary of section names to content
        'contact_info': contact_info, # Dictionary of contact details lists
        'entities': entities, # Dictionary of extracted entities
        'completeness_score': completeness_score, # Calculated score
        'text': text # Include the raw extracted text for reference
    }

# Note: The semantic_matcher.py module would need to use the 'text' and 'skills'
# from the output of this analyze_resume function.
# extract_resume_keywords and extract_skills_from_resume might be redundant
# if analyze_resume_skills already provides them.
# Ensure analyze_resume_skills is separate and focuses just on skills/keywords.

# Example of how analyze_resume_skills might look (assuming it exists)
# def analyze_resume_skills(text):
#     # Your logic to extract skills and keywords
#     skills = ["skill1", "skill2"] # Example
#     keywords = ["keyword1", "keyword2"] # Example
#     return {"skills": skills, "keywords": keywords}
